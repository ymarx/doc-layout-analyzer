#!/usr/bin/env python3
"""
System Test Script
CPU/GPU 듀얼 모드 시스템 테스트
"""

import sys
import asyncio
from pathlib import Path
import tempfile
import json
import time

# 프로젝트 경로 추가
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from src.core.config import ConfigManager
from src.core.device_manager import DeviceManager
from src.parsers.docx_parser import DocxParser
from src.parsers.pdf_parser import PDFParser
from src.core.docjson import DocJSONConverter
from src.analyzers.layout_analyzer import LayoutAnalyzer
from src.extractors.content_extractor import ContentExtractor

print("🧪 Document Layout Analyzer - System Test")
print("="*50)

def test_imports():
    """필수 라이브러리 임포트 테스트"""
    print("\n1. 📦 라이브러리 임포트 테스트")

    required_libs = [
        ("yaml", "PyYAML"),
        ("PIL", "Pillow"),
        ("numpy", "NumPy"),
        ("pandas", "Pandas"),
    ]

    optional_libs = [
        ("paddleocr", "PaddleOCR"),
        ("docx", "python-docx"),
        ("fitz", "PyMuPDF"),
        ("pdfplumber", "pdfplumber"),
        ("camelot", "camelot-py"),
        ("torch", "PyTorch"),
    ]

    print("  필수 라이브러리:")
    for lib, name in required_libs:
        try:
            __import__(lib)
            print(f"  ✅ {name}")
        except ImportError:
            print(f"  ❌ {name} - 설치 필요")

    print("  선택적 라이브러리:")
    for lib, name in optional_libs:
        try:
            __import__(lib)
            print(f"  ✅ {name}")
        except ImportError:
            print(f"  ⚠️  {name} - 설치되지 않음 (기능 제한)")

def test_config_system():
    """설정 시스템 테스트"""
    print("\n2. ⚙️ 설정 시스템 테스트")

    try:
        config = ConfigManager()
        print(f"  ✅ 설정 로드 완료")
        print(f"  - 처리 모드: {config.system.processing_mode.value}")
        print(f"  - 최대 워커: {config.system.max_workers}")
        print(f"  - OCR GPU 사용: {config.ocr.use_gpu}")
        print(f"  - 임베딩 디바이스: {config.embedding.device}")

    except Exception as e:
        print(f"  ❌ 설정 시스템 오류: {e}")

def test_device_manager():
    """디바이스 매니저 테스트"""
    print("\n3. 🖥️ 디바이스 매니저 테스트")

    try:
        dm = DeviceManager()
        print(f"  ✅ 디바이스 매니저 초기화 완료")
        print(f"  - 감지된 디바이스: {len(dm.devices)}개")
        print(f"  - 현재 디바이스: {dm.current_device}")

        # 벤치마크 테스트
        print("  🏃 성능 벤치마크 (간단한 연산):")
        for device in dm.devices[:2]:  # 최대 2개만 테스트
            if device.is_available:
                device_name = f"{device.device_type.value}:{device.device_id}" if device.device_type.value != "cpu" else "cpu"
                result = dm.benchmark_device(device_name, iterations=50)
                if "error" not in result:
                    print(f"    - {device_name}: {result['gflops']:.2f} GFLOPS")
                else:
                    print(f"    - {device_name}: 벤치마크 실패")

    except Exception as e:
        print(f"  ❌ 디바이스 매니저 오류: {e}")

def test_docjson_converter():
    """DocJSON 변환기 테스트"""
    print("\n4. 📄 DocJSON 변환기 테스트")

    try:
        converter = DocJSONConverter()
        print("  ✅ DocJSON 변환기 초기화 완료")

        # 스키마 유효성 검사
        test_doc = {
            "version": "2.0",
            "doc_id": "test-uuid-12345",
            "metadata": {
                "title": "테스트 문서",
                "doc_type": "기술기준",
                "language": ["ko", "en"],
                "created": "2025-01-01T00:00:00",
                "pages": 1,
                "file_size": 1024
            },
            "sections": []
        }

        if converter.validate_docjson(test_doc):
            print("  ✅ 스키마 검증 통과")
        else:
            print("  ❌ 스키마 검증 실패")

    except Exception as e:
        print(f"  ❌ DocJSON 변환기 오류: {e}")

async def test_parsers():
    """파서 테스트"""
    print("\n5. 📖 파서 테스트")

    try:
        # DOCX 파서 테스트
        print("  DOCX 파서:")
        docx_parser = DocxParser()
        print("    ✅ DOCX 파서 초기화 완료")
        print(f"    - 지원 형식: {docx_parser.supported_formats}")

        # PDF 파서 테스트
        print("  PDF 파서:")
        pdf_parser = PDFParser()
        print("    ✅ PDF 파서 초기화 완료")
        print(f"    - 지원 형식: {pdf_parser.supported_formats}")

        # 샘플 텍스트로 타입 감지 테스트
        test_files = [
            "test.docx",
            "document.pdf",
            "image.png",
            "unknown.xyz"
        ]

        print("  파일 타입 감지 테스트:")
        for filename in test_files:
            doc_type = docx_parser.detect_document_type(filename)
            print(f"    - {filename}: {doc_type.value}")

    except Exception as e:
        print(f"  ❌ 파서 테스트 오류: {e}")

async def test_layout_analyzer():
    """레이아웃 분석기 테스트"""
    print("\n6. 🔍 레이아웃 분석기 테스트")

    try:
        dm = DeviceManager()
        config = ConfigManager()
        analyzer = LayoutAnalyzer(dm, config)

        print("  ✅ 레이아웃 분석기 초기화 완료")
        print(f"  - 사용 가능한 분석기: {analyzer.available_analyzers}")
        print(f"  - 현재 디바이스: {analyzer.current_device}")

    except Exception as e:
        print(f"  ❌ 레이아웃 분석기 오류: {e}")

async def test_content_extractor():
    """콘텐츠 추출기 테스트"""
    print("\n7. 📝 콘텐츠 추출기 테스트")

    try:
        dm = DeviceManager()
        config = ConfigManager()
        extractor = ContentExtractor(dm, config)

        print("  ✅ 콘텐츠 추출기 초기화 완료")
        print(f"  - 현재 디바이스: {extractor.current_device}")

    except Exception as e:
        print(f"  ❌ 콘텐츠 추출기 오류: {e}")

def create_sample_docx():
    """샘플 DOCX 파일 생성"""
    try:
        from docx import Document

        doc = Document()
        doc.add_heading('테스트 문서', 0)
        doc.add_heading('1. 서론', level=1)
        doc.add_paragraph('이것은 테스트를 위한 샘플 문서입니다.')

        # 표 추가
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '항목'
        table.cell(0, 1).text = '값'
        table.cell(1, 0).text = '온도'
        table.cell(1, 1).text = '1500℃'

        doc.add_heading('2. 결론', level=1)
        doc.add_paragraph('테스트 문서 작성이 완료되었습니다.')

        # 임시 파일로 저장
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(temp_file.name)

        return temp_file.name

    except ImportError:
        print("  ⚠️  python-docx가 설치되지 않아 DOCX 샘플 생성을 건너뜁니다.")
        return None
    except Exception as e:
        print(f"  ❌ 샘플 DOCX 생성 실패: {e}")
        return None

async def test_full_pipeline():
    """전체 파이프라인 테스트"""
    print("\n8. 🔄 전체 파이프라인 테스트")

    # 샘플 파일 생성
    sample_file = create_sample_docx()
    if not sample_file:
        print("  ⚠️  샘플 파일 생성 실패, 파이프라인 테스트를 건너뜁니다.")
        return

    try:
        print(f"  📄 샘플 파일 생성: {sample_file}")

        # 메인 애플리케이션 컴포넌트 초기화
        from main import DocumentAnalyzer

        analyzer = DocumentAnalyzer()
        print("  ✅ Document Analyzer 초기화 완료")

        # 문서 처리 테스트
        print("  🔄 문서 처리 시작...")
        start_time = time.time()

        # 출력 디렉토리 생성
        output_dir = Path("./test_output")
        output_dir.mkdir(exist_ok=True)

        success = await analyzer.process_document(sample_file, str(output_dir))

        processing_time = time.time() - start_time

        if success:
            print(f"  ✅ 파이프라인 테스트 성공!")
            print(f"  - 처리 시간: {processing_time:.2f}초")

            # 결과 파일 확인
            output_file = output_dir / (Path(sample_file).stem + ".docjson")
            if output_file.exists():
                print(f"  📊 결과 파일 크기: {output_file.stat().st_size:,} bytes")

                # JSON 유효성 확인
                try:
                    with open(output_file, 'r', encoding='utf-8') as f:
                        result_data = json.load(f)
                    print(f"  📄 생성된 섹션: {len(result_data.get('sections', []))}개")
                    print(f"  📋 전체 블록: {sum(len(section.get('blocks', [])) for section in result_data.get('sections', []))}개")
                except Exception as e:
                    print(f"  ⚠️  결과 파일 검증 실패: {e}")
            else:
                print("  ⚠️  결과 파일을 찾을 수 없습니다.")

        else:
            print("  ❌ 파이프라인 테스트 실패")

    except Exception as e:
        print(f"  ❌ 파이프라인 테스트 오류: {e}")

    finally:
        # 임시 파일 정리
        try:
            Path(sample_file).unlink(missing_ok=True)
        except:
            pass

def print_summary():
    """테스트 요약 출력"""
    print("\n" + "="*50)
    print("📊 테스트 완료 요약")
    print("="*50)
    print("✅ 성공한 테스트는 정상 동작합니다.")
    print("❌ 실패한 테스트는 해당 기능에 문제가 있습니다.")
    print("⚠️  경고가 있는 테스트는 일부 기능이 제한됩니다.")
    print()
    print("🚀 시스템 사용 준비:")
    print("  python main.py --info  # 시스템 정보 확인")
    print("  python main.py document.pdf  # 문서 처리")
    print()

async def main():
    """메인 테스트 함수"""
    try:
        print("CPU/GPU 듀얼 모드 지원 확인 및 시스템 테스트를 시작합니다.\n")

        # 단계별 테스트 실행
        test_imports()
        test_config_system()
        test_device_manager()
        test_docjson_converter()
        await test_parsers()
        await test_layout_analyzer()
        await test_content_extractor()
        await test_full_pipeline()

        print_summary()

    except KeyboardInterrupt:
        print("\n테스트가 사용자에 의해 중단되었습니다.")
    except Exception as e:
        print(f"\n❌ 테스트 실행 중 오류: {e}")

if __name__ == "__main__":
    asyncio.run(main())