"""
DOCX Parser - Microsoft Word 문서 파서
python-docx 기반으로 텍스트, 표, 이미지 추출
"""

import logging
import time
import asyncio
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, BinaryIO
import io
import zipfile

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shape import InlineShape
from docx.enum.shape import WD_INLINE_SHAPE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

from .base_parser import BaseParser, ParseResult, ProcessingOptions, DocumentType

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """DOCX 파서"""

    def __init__(self, device_manager=None, config=None):
        super().__init__(device_manager, config)
        self.supported_formats = ['.docx']

    def can_handle(self, file_path: Union[str, Path]) -> bool:
        """DOCX 파일 처리 가능 여부"""
        if isinstance(file_path, (str, Path)):
            return Path(file_path).suffix.lower() == '.docx'
        return False

    async def parse(self,
                   file_path: Union[str, Path, BinaryIO],
                   options: ProcessingOptions = None) -> ParseResult:
        """DOCX 문서 파싱"""
        start_time = time.time()
        options = options or ProcessingOptions()

        try:
            # 파일 유효성 검사
            if isinstance(file_path, (str, Path)) and not self.validate_file(file_path):
                return ParseResult(
                    success=False,
                    document_type=DocumentType.DOCX,
                    error="File validation failed"
                )

            # 문서 로드
            if isinstance(file_path, (str, Path)):
                doc = Document(str(file_path))
                file_info = self.get_file_info(file_path)
            else:
                doc = Document(file_path)
                file_info = {"filename": "stream_input", "size": 0}

            # 문서 내용 추출
            content = await self._extract_content(doc, options)

            # 메타데이터 추출
            metadata = self._extract_metadata(doc, file_info)

            processing_time = time.time() - start_time
            self.update_stats(processing_time, True)

            return ParseResult(
                success=True,
                document_type=DocumentType.DOCX,
                content=content,
                metadata=metadata,
                processing_time=processing_time,
                pages=self._estimate_pages(doc),
                file_size=file_info.get("size", 0)
            )

        except Exception as e:
            processing_time = time.time() - start_time
            self.update_stats(processing_time, False)
            logger.error(f"DOCX 파싱 실패: {e}")

            return ParseResult(
                success=False,
                document_type=DocumentType.DOCX,
                error=str(e),
                processing_time=processing_time
            )

    async def _extract_content(self, doc: DocxDocument, options: ProcessingOptions) -> Dict[str, Any]:
        """문서 내용 추출"""
        content = {
            "sections": [],
            "paragraphs": [],
            "tables": [],
            "images": [],
            "headers_footers": {}
        }

        # 텍스트 추출
        if options.extract_text:
            content["paragraphs"] = await self._extract_paragraphs(doc)

        # 표 추출
        if options.extract_tables:
            content["tables"] = await self._extract_tables(doc)

        # 이미지 추출
        if options.extract_images:
            content["images"] = await self._extract_images(doc)

        # 헤더/푸터 추출
        content["headers_footers"] = await self._extract_headers_footers(doc)

        # 섹션 구조화
        content["sections"] = self._organize_into_sections(content)

        return content

    async def _extract_paragraphs(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """단락 추출"""
        paragraphs = []

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # 빈 단락 제외
                para_data = {
                    "id": f"para_{i}",
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "alignment": str(para.alignment) if para.alignment else None,
                    "runs": []
                }

                # Run 정보 추출 (글꼴, 색상 등)
                for j, run in enumerate(para.runs):
                    run_data = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name,
                        "font_size": run.font.size.pt if run.font.size else None
                    }
                    para_data["runs"].append(run_data)

                paragraphs.append(para_data)

        return paragraphs

    async def _extract_tables(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """표 추출"""
        tables = []

        for i, table in enumerate(doc.tables):
            table_data = {
                "id": f"table_{i}",
                "rows": len(table.rows),
                "cols": len(table.columns) if table.columns else 0,
                "data": [],
                "style": table.style.name if table.style else None,
                "alignment": str(table.alignment) if hasattr(table, 'alignment') else None
            }

            # 셀 데이터 추출
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    cell_data = {
                        "text": cell.text.strip(),
                        "row_span": 1,  # python-docx에서 병합 셀 정보는 복잡함
                        "col_span": 1,
                        "paragraphs": [p.text for p in cell.paragraphs if p.text.strip()]
                    }
                    row_data.append(cell_data)
                table_data["data"].append(row_data)

            # 표 캡션 찾기 (표 앞/뒤 단락에서)
            table_data["caption"] = self._find_table_caption(doc, i)

            tables.append(table_data)

        return tables

    async def _extract_images(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """이미지 추출"""
        images = []

        # 문서 내 인라인 이미지 추출
        for i, paragraph in enumerate(doc.paragraphs):
            for j, run in enumerate(paragraph.runs):
                for k, inline_shape in enumerate(run.element.xpath('.//w:drawing')):
                    try:
                        image_data = {
                            "id": f"image_{i}_{j}_{k}",
                            "type": "inline",
                            "paragraph_id": f"para_{i}",
                            "caption": self._find_image_caption(doc, i),
                            "alt_text": "",  # DOCX에서 alt text 추출은 복잡
                            "width": None,
                            "height": None
                        }
                        images.append(image_data)
                    except Exception as e:
                        logger.debug(f"이미지 추출 실패: {e}")

        return images

    async def _extract_headers_footers(self, doc: DocxDocument) -> Dict[str, Any]:
        """헤더/푸터 추출"""
        headers_footers = {
            "headers": [],
            "footers": []
        }

        try:
            # 섹션별 헤더/푸터
            for section in doc.sections:
                # 헤더
                if section.header:
                    header_text = []
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            header_text.append(para.text)
                    if header_text:
                        headers_footers["headers"].append({
                            "text": "\n".join(header_text),
                            "type": "default"
                        })

                # 푸터
                if section.footer:
                    footer_text = []
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            footer_text.append(para.text)
                    if footer_text:
                        headers_footers["footers"].append({
                            "text": "\n".join(footer_text),
                            "type": "default"
                        })

        except Exception as e:
            logger.debug(f"헤더/푸터 추출 실패: {e}")

        return headers_footers

    def _extract_metadata(self, doc: DocxDocument, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """메타데이터 추출"""
        metadata = file_info.copy()

        try:
            core_props = doc.core_properties

            metadata.update({
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "last_modified_by": core_props.last_modified_by,
                "revision": core_props.revision,
                "language": core_props.language,
                "category": core_props.category,
                "comments": core_props.comments
            })

        except Exception as e:
            logger.debug(f"메타데이터 추출 실패: {e}")

        return metadata

    def _estimate_pages(self, doc: DocxDocument) -> int:
        """페이지 수 추정"""
        try:
            # 단락 수 기반 대략적 추정
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            table_count = len(doc.tables)

            # 단락 40개당 1페이지, 표 2개당 1페이지로 추정
            estimated_pages = max(1, (paragraph_count // 40) + (table_count // 2) + 1)
            return estimated_pages

        except Exception:
            return 1

    def _find_table_caption(self, doc: DocxDocument, table_index: int) -> Optional[str]:
        """표 캡션 찾기"""
        try:
            # 표 앞뒤 단락에서 '표', 'Table' 키워드 찾기
            all_elements = []

            # 모든 요소를 순서대로 수집 (단락과 표)
            for element in doc.element.body:
                if element.tag.endswith('tbl'):
                    all_elements.append(('table', element))
                elif element.tag.endswith('p'):
                    all_elements.append(('paragraph', element))

            # 현재 표의 위치 찾기
            table_count = 0
            for i, (elem_type, elem) in enumerate(all_elements):
                if elem_type == 'table':
                    if table_count == table_index:
                        # 앞뒤 단락 확인
                        for j in range(max(0, i-2), min(len(all_elements), i+3)):
                            if j != i and all_elements[j][0] == 'paragraph':
                                para_text = all_elements[j][1].text or ""
                                if any(keyword in para_text.lower()
                                      for keyword in ['표', 'table', '그림', 'figure']):
                                    return para_text.strip()
                        break
                    table_count += 1

        except Exception as e:
            logger.debug(f"표 캡션 찾기 실패: {e}")

        return None

    def _find_image_caption(self, doc: DocxDocument, paragraph_index: int) -> Optional[str]:
        """이미지 캡션 찾기"""
        try:
            # 이미지가 있는 단락 앞뒤에서 캡션 찾기
            for i in range(max(0, paragraph_index-2),
                          min(len(doc.paragraphs), paragraph_index+3)):
                if i != paragraph_index:
                    para_text = doc.paragraphs[i].text
                    if any(keyword in para_text.lower()
                          for keyword in ['그림', 'figure', '이미지', 'image']):
                        return para_text.strip()

        except Exception as e:
            logger.debug(f"이미지 캡션 찾기 실패: {e}")

        return None

    def _organize_into_sections(self, content: Dict[str, Any]) -> List[Dict[str, Any]]:
        """내용을 섹션으로 구조화"""
        sections = []
        current_section = None

        for para in content["paragraphs"]:
            # 제목 스타일 확인
            style = para.get("style", "").lower()
            text = para.get("text", "")

            # 제목 판단 (스타일 기반)
            is_heading = any(heading in style for heading in
                           ['heading', '제목', 'title']) or self._is_heading_text(text)

            if is_heading:
                # 새 섹션 시작
                if current_section:
                    sections.append(current_section)

                current_section = {
                    "id": f"section_{len(sections)}",
                    "heading": text,
                    "level": self._get_heading_level(style),
                    "paragraphs": [],
                    "tables": [],
                    "images": []
                }
            else:
                # 현재 섹션에 단락 추가
                if current_section is None:
                    current_section = {
                        "id": "section_0",
                        "heading": "서론",
                        "level": 1,
                        "paragraphs": [],
                        "tables": [],
                        "images": []
                    }

                current_section["paragraphs"].append(para)

        # 마지막 섹션 추가
        if current_section:
            sections.append(current_section)

        # 표와 이미지를 적절한 섹션에 배치
        self._assign_tables_and_images_to_sections(sections, content["tables"], content["images"])

        return sections

    def _is_heading_text(self, text: str) -> bool:
        """텍스트가 제목인지 판단"""
        if not text or len(text) > 100:  # 너무 긴 텍스트는 제목이 아님
            return False

        # 번호가 있는 제목 패턴 (1. 1.1 가. 가) 등)
        import re
        heading_patterns = [
            r'^\d+\.\s+',  # 1. 제목
            r'^\d+\.\d+\s+',  # 1.1 제목
            r'^[가-힣]\.\s+',  # 가. 제목
            r'^[가-힣]\)\s+',  # 가) 제목
            r'^\([가-힣]\)\s+',  # (가) 제목
            r'^\d+\)\s+',  # 1) 제목
        ]

        for pattern in heading_patterns:
            if re.match(pattern, text):
                return True

        return False

    def _get_heading_level(self, style: str) -> int:
        """제목 레벨 추출"""
        style_lower = style.lower()

        if 'heading 1' in style_lower or '제목 1' in style_lower:
            return 1
        elif 'heading 2' in style_lower or '제목 2' in style_lower:
            return 2
        elif 'heading 3' in style_lower or '제목 3' in style_lower:
            return 3
        elif 'heading 4' in style_lower or '제목 4' in style_lower:
            return 4
        elif 'heading 5' in style_lower or '제목 5' in style_lower:
            return 5
        elif 'heading 6' in style_lower or '제목 6' in style_lower:
            return 6

        return 1  # 기본값

    def _assign_tables_and_images_to_sections(self, sections: List[Dict[str, Any]],
                                            tables: List[Dict[str, Any]],
                                            images: List[Dict[str, Any]]):
        """표와 이미지를 섹션에 배치"""
        # 간단한 구현: 모든 표와 이미지를 첫 번째 섹션에 배치
        # 더 정교한 구현을 위해서는 XML 순서 분석 필요
        if sections:
            sections[0]["tables"] = tables
            sections[0]["images"] = images