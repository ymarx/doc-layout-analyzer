"""
Unified DOCX Parser - 통합 DOCX 파서
기존 DocxParser와 DocxEnhancedParser 기능을 통합한 단일 파서
"""

import logging
import time
import asyncio
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, BinaryIO, Tuple
from dataclasses import dataclass, field
import re
from datetime import datetime

# python-docx imports
from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shape import InlineShape
from docx.enum.shape import WD_INLINE_SHAPE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

# Base imports
from .base_parser import BaseParser, ParseResult, ProcessingOptions, DocumentType
from ..core.docjson import DocJSON, BoundingBox, ContentBlock, SemanticInfo, ContentBlockType, DocumentMetadata, DocumentSection
from .document_analyzer import DocumentAnalyzer
from .diagram_flow_analyzer import DiagramFlowAnalyzer

logger = logging.getLogger(__name__)

# DOCX XML 네임스페이스 정의
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    've': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'o': 'urn:schemas-microsoft-com:office:office',
    'v': 'urn:schemas-microsoft-com:vml',
    'w10': 'urn:schemas-microsoft-com:office:word',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
    'wpi': 'http://schemas.microsoft.com/office/word/2010/wordprocessingInk',
    'wne': 'http://schemas.microsoft.com/office/word/2006/wordml',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
}


@dataclass
class ParsingMode:
    """파싱 모드 설정"""
    use_python_docx: bool = True      # python-docx 사용 여부
    use_xml_analysis: bool = True     # XML 직접 분석 여부
    extract_diagrams: bool = True     # 다이어그램 추출
    deep_structure: bool = True       # 깊은 구조 분석
    include_metadata: bool = True     # 메타데이터 포함


class UnifiedDocxParser(BaseParser):
    """통합 DOCX 파서 - 기본 파서와 고급 파서 기능 통합"""

    def __init__(self, device_manager=None, config=None):
        super().__init__(device_manager, config)
        self.supported_formats = ['.docx']

        # 파싱 모드 설정
        self.parsing_modes = {
            'basic': ParsingMode(
                use_python_docx=True,
                use_xml_analysis=False,
                extract_diagrams=False,
                deep_structure=False
            ),
            'enhanced': ParsingMode(
                use_python_docx=True,
                use_xml_analysis=True,
                extract_diagrams=True,
                deep_structure=True
            ),
            'xml_only': ParsingMode(
                use_python_docx=False,
                use_xml_analysis=True,
                extract_diagrams=True,
                deep_structure=True
            )
        }

    def can_handle(self, file_path: Union[str, Path]) -> bool:
        """DOCX 파일 처리 가능 여부"""
        if isinstance(file_path, (str, Path)):
            return Path(file_path).suffix.lower() == '.docx'
        return False

    async def parse(self,
                   file_path: Union[str, Path, BinaryIO],
                   options: ProcessingOptions = None) -> ParseResult:
        """통합 DOCX 문서 파싱"""
        start_time = time.time()
        options = options or ProcessingOptions()

        try:
            # 파일 유효성 검사
            if isinstance(file_path, (str, Path)) and not self.validate_file(file_path):
                return ParseResult(
                    success=False,
                    document_type=DocumentType.DOCX,
                    error="Invalid DOCX file"
                )

            # 파싱 모드 결정
            parsing_mode = self._determine_parsing_mode(options)

            # 병렬 파싱 실행
            docx_content = None
            xml_content = None

            tasks = []

            if parsing_mode.use_python_docx:
                tasks.append(self._parse_with_python_docx(file_path, parsing_mode))

            if parsing_mode.use_xml_analysis:
                tasks.append(self._parse_with_xml_analysis(file_path, parsing_mode))

            # 병렬 실행
            if tasks:
                results = await asyncio.gather(*tasks, return_exceptions=True)

                # 결과 처리
                for i, result in enumerate(results):
                    if not isinstance(result, Exception):
                        if i == 0 and parsing_mode.use_python_docx:
                            docx_content = result
                        elif (i == 1 and parsing_mode.use_python_docx and parsing_mode.use_xml_analysis) or \
                             (i == 0 and not parsing_mode.use_python_docx and parsing_mode.use_xml_analysis):
                            xml_content = result

            # 결과 통합
            combined_content = self._combine_parsing_results(docx_content, xml_content, parsing_mode)

            # 고급 문서 구조 분석 추가
            analyzer = DocumentAnalyzer()
            document_structure = analyzer.analyze_document(combined_content)
            recognition_score = analyzer.get_recognition_score(document_structure)

            # 구조 분석 결과를 combined_content에 추가
            combined_content['document_structure'] = {
                'document_number': document_structure.document_number,
                'title': document_structure.title,
                'author': document_structure.author,
                'effective_date': document_structure.effective_date,
                'revision': document_structure.revision,
                'sections': document_structure.sections,
                'patterns_found': document_structure.patterns_found,
                'metadata': document_structure.metadata,
                'recognition_score': recognition_score
            }

            # DocJSON 변환 (구조 분석 결과 포함)
            docjson = self._convert_to_docjson(combined_content, file_path)

            processing_time = time.time() - start_time

            return ParseResult(
                success=True,
                document_type=DocumentType.DOCX,
                content={
                    'docjson': docjson.to_dict() if docjson else None,
                    'raw_content': combined_content,
                    'parsing_mode': parsing_mode.__dict__,
                    'document_structure': combined_content.get('document_structure', {})
                },
                metadata={
                    'parsing_method': 'unified',
                    'python_docx_used': parsing_mode.use_python_docx,
                    'xml_analysis_used': parsing_mode.use_xml_analysis,
                    'processing_time': processing_time,
                    'parser_version': '2.0',
                    'recognition_score': recognition_score
                },
                processing_time=processing_time,
                pages=combined_content.get('page_count', 1) if combined_content else 1
            )

        except Exception as e:
            logger.error(f"DOCX 파싱 실패: {e}")
            return ParseResult(
                success=False,
                document_type=DocumentType.DOCX,
                error=str(e),
                processing_time=time.time() - start_time
            )

    def _determine_parsing_mode(self, options: ProcessingOptions) -> ParsingMode:
        """처리 옵션에 따른 파싱 모드 결정"""
        # 옵션에 따른 모드 선택 로직
        if hasattr(options, 'parsing_complexity'):
            if options.parsing_complexity == 'basic':
                return self.parsing_modes['basic']
            elif options.parsing_complexity == 'enhanced':
                return self.parsing_modes['enhanced']
            elif options.parsing_complexity == 'xml_only':
                return self.parsing_modes['xml_only']

        # 기본값: enhanced 모드
        return self.parsing_modes['enhanced']

    async def _parse_with_python_docx(self, file_path: Union[str, Path, BinaryIO], mode: ParsingMode) -> Dict[str, Any]:
        """python-docx를 이용한 파싱"""
        try:
            # 기존 DocxParser 로직 통합
            if isinstance(file_path, (str, Path)):
                doc = Document(str(file_path))
            else:
                doc = Document(file_path)

            content = {
                'method': 'python_docx',
                'paragraphs': [],
                'tables': [],
                'images': [],
                'page_count': 1
            }

            # 단락 처리
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    content['paragraphs'].append({
                        'index': i,
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal',
                        'alignment': str(para.alignment) if para.alignment else None
                    })

            # 표 처리
            for i, table in enumerate(doc.tables):
                table_data = {
                    'index': i,
                    'rows': len(table.rows),
                    'cols': len(table.columns) if table.rows else 0,
                    'data': []
                }

                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data['data'].append(row_data)

                content['tables'].append(table_data)

            # 이미지 처리 (간단화)
            for i, para in enumerate(doc.paragraphs):
                for run in para.runs:
                    for inline_shape in run.element.xpath('.//wp:inline'):
                        content['images'].append({
                            'paragraph_index': i,
                            'type': 'inline_image'
                        })

            return content

        except Exception as e:
            logger.error(f"python-docx 파싱 실패: {e}")
            return {'method': 'python_docx', 'error': str(e)}

    async def _parse_with_xml_analysis(self, file_path: Union[str, Path, BinaryIO], mode: ParsingMode) -> Dict[str, Any]:
        """XML 직접 분석을 통한 파싱"""
        try:
            # 기존 DocxEnhancedParser 로직 통합
            content = {
                'method': 'xml_analysis',
                'structure': {},
                'diagrams': [],
                'advanced_elements': []
            }

            # ZIP 파일로 DOCX 열기
            if isinstance(file_path, (str, Path)):
                with zipfile.ZipFile(file_path, 'r') as docx_zip:
                    return await self._analyze_docx_xml(docx_zip, mode)
            else:
                # BinaryIO 처리
                with zipfile.ZipFile(file_path, 'r') as docx_zip:
                    return await self._analyze_docx_xml(docx_zip, mode)

        except Exception as e:
            logger.error(f"XML 분석 실패: {e}")
            return {'method': 'xml_analysis', 'error': str(e)}

    async def _analyze_docx_xml(self, docx_zip: zipfile.ZipFile, mode: ParsingMode) -> Dict[str, Any]:
        """DOCX XML 구조 분석"""
        content = {
            'method': 'xml_analysis',
            'structure': {},
            'diagrams': [],
            'advanced_elements': [],
            'relationships': {},
            'process_flows': []  # 프로세스 흐름 추가
        }

        try:
            # document.xml 분석
            if 'word/document.xml' in docx_zip.namelist():
                document_xml = docx_zip.read('word/document.xml')
                root = ET.fromstring(document_xml)

                # 네임스페이스 등록
                for prefix, uri in NAMESPACES.items():
                    ET.register_namespace(prefix, uri)

                # 구조 분석
                content['structure'] = self._analyze_document_structure(root)

                # 다이어그램 추출 (mode.extract_diagrams가 True인 경우)
                if mode.extract_diagrams:
                    content['diagrams'] = self._extract_diagrams(root, docx_zip)

                    # 다이어그램에서 프로세스 흐름 추출
                    diagram_analyzer = DiagramFlowAnalyzer()
                    process_flows = diagram_analyzer.analyze_diagrams(docx_zip, root)
                    content['process_flows'] = diagram_analyzer.format_flow_for_docjson(process_flows)

            # 헤더/푸터 분석 추가
            headers_footers = self._extract_headers_footers(docx_zip)
            if headers_footers:
                content['headers_footers'] = headers_footers

            # 관계 파일 분석
            if 'word/_rels/document.xml.rels' in docx_zip.namelist():
                rels_xml = docx_zip.read('word/_rels/document.xml.rels')
                content['relationships'] = self._analyze_relationships(rels_xml)

        except Exception as e:
            logger.error(f"XML 구조 분석 실패: {e}")
            content['error'] = str(e)

        return content

    def _analyze_document_structure(self, root: ET.Element) -> Dict[str, Any]:
        """문서 구조 분석"""
        structure = {
            'paragraphs': [],
            'tables': [],
            'sections': []
        }

        # 단락 분석
        for para in root.findall('.//w:p', NAMESPACES):
            para_text = ''.join(t.text for t in para.findall('.//w:t', NAMESPACES) if t.text)
            if para_text.strip():
                structure['paragraphs'].append({
                    'text': para_text,
                    'style_info': self._extract_paragraph_style(para)
                })

        # 표 분석
        for table in root.findall('.//w:tbl', NAMESPACES):
            table_data = self._extract_table_data(table)
            structure['tables'].append(table_data)

        return structure

    def _extract_paragraph_style(self, para: ET.Element) -> Dict[str, Any]:
        """단락 스타일 정보 추출"""
        style_info = {}

        # 스타일 참조
        style_ref = para.find('.//w:pStyle', NAMESPACES)
        if style_ref is not None:
            style_info['style_name'] = style_ref.get(f'{{{NAMESPACES["w"]}}}val')

        return style_info

    def _extract_table_data(self, table: ET.Element) -> Dict[str, Any]:
        """표 데이터 추출"""
        table_data = {
            'rows': [],
            'properties': {}
        }

        for row in table.findall('.//w:tr', NAMESPACES):
            row_data = []
            for cell in row.findall('.//w:tc', NAMESPACES):
                cell_text = ''.join(t.text for t in cell.findall('.//w:t', NAMESPACES) if t.text)
                row_data.append(cell_text.strip())
            table_data['rows'].append(row_data)

        return table_data

    def _extract_diagrams(self, root: ET.Element, docx_zip: zipfile.ZipFile) -> List[Dict[str, Any]]:
        """다이어그램 추출"""
        diagrams = []

        # Drawing 요소 찾기
        for drawing in root.findall('.//w:drawing', NAMESPACES):
            diagram_info = {
                'type': 'drawing',
                'description': 'Embedded diagram or shape'
            }

            # 추가 다이어그램 정보 추출 로직
            diagrams.append(diagram_info)

        return diagrams

    def _analyze_relationships(self, rels_xml: bytes) -> Dict[str, Any]:
        """관계 파일 분석"""
        relationships = {}

        try:
            root = ET.fromstring(rels_xml)
            for rel in root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                rel_id = rel.get('Id')
                rel_type = rel.get('Type')
                rel_target = rel.get('Target')

                relationships[rel_id] = {
                    'type': rel_type,
                    'target': rel_target
                }
        except Exception as e:
            logger.error(f"관계 분석 실패: {e}")

        return relationships

    def _extract_headers_footers(self, docx_zip: zipfile.ZipFile) -> Dict[str, Any]:
        """헤더와 푸터 내용 추출"""
        headers_footers = {
            'headers': [],
            'footers': []
        }

        try:
            # 헤더 파일들 찾기
            for file_name in docx_zip.namelist():
                if 'header' in file_name and file_name.endswith('.xml'):
                    try:
                        header_xml = docx_zip.read(file_name)
                        root = ET.fromstring(header_xml)

                        # 텍스트 추출
                        texts = []
                        for t in root.findall('.//w:t', NAMESPACES):
                            if t.text:
                                texts.append(t.text)

                        if texts:
                            headers_footers['headers'].append({
                                'file': file_name,
                                'text': ' '.join(texts)
                            })

                    except Exception as e:
                        logger.warning(f"헤더 {file_name} 읽기 실패: {e}")

                elif 'footer' in file_name and file_name.endswith('.xml'):
                    try:
                        footer_xml = docx_zip.read(file_name)
                        root = ET.fromstring(footer_xml)

                        # 텍스트 추출
                        texts = []
                        for t in root.findall('.//w:t', NAMESPACES):
                            if t.text:
                                texts.append(t.text)

                        if texts:
                            headers_footers['footers'].append({
                                'file': file_name,
                                'text': ' '.join(texts)
                            })

                    except Exception as e:
                        logger.warning(f"푸터 {file_name} 읽기 실패: {e}")

        except Exception as e:
            logger.error(f"헤더/푸터 추출 실패: {e}")

        return headers_footers

    def _combine_parsing_results(self, docx_content: Dict[str, Any], xml_content: Dict[str, Any], mode: ParsingMode) -> Dict[str, Any]:
        """두 파싱 결과를 통합"""
        combined = {
            'combined': True,
            'sources': []
        }

        if docx_content and 'error' not in docx_content:
            combined.update(docx_content)
            combined['sources'].append('python_docx')

        if xml_content and 'error' not in xml_content:
            # XML 분석 결과 통합
            if 'structure' in xml_content:
                combined['xml_structure'] = xml_content['structure']
            if 'diagrams' in xml_content:
                combined['diagrams'] = xml_content['diagrams']
            if 'relationships' in xml_content:
                combined['relationships'] = xml_content['relationships']
            if 'headers_footers' in xml_content:
                combined['headers_footers'] = xml_content['headers_footers']
                # XML 구조에도 추가
                if 'xml_structure' not in combined:
                    combined['xml_structure'] = {}
                combined['xml_structure']['headers_footers'] = xml_content['headers_footers']
            if 'process_flows' in xml_content:
                combined['process_flows'] = xml_content['process_flows']
            combined['sources'].append('xml_analysis')

        return combined

    def _convert_to_docjson(self, content: Dict[str, Any], file_path: Union[str, Path, BinaryIO]) -> Optional[DocJSON]:
        """파싱 결과를 DocJSON으로 변환"""
        if not content or 'error' in content:
            return None

        try:
            # 구조 분석 결과 활용
            structure = content.get('document_structure', {})

            # 메타데이터 생성 (구조 분석 결과 우선 사용)
            file_name = str(file_path) if isinstance(file_path, (str, Path)) else "document.docx"

            # 문서 제목: 구조 분석에서 찾은 제목 우선 사용
            doc_title = structure.get('title') or Path(file_name).stem

            # 작성자: 구조 분석에서 찾은 작성자 우선 사용
            doc_author = structure.get('author') or "Unknown"

            metadata = DocumentMetadata(
                title=doc_title,
                doc_type="docx",
                author=doc_author,
                created=datetime.now().isoformat(),
                modified=datetime.now().isoformat(),
                pages=content.get('page_count', 1),
                file_size=0,
                # 구조 분석에서 추출한 추가 정보
                document_number=structure.get('document_number'),
                revision=structure.get('revision'),
                effective_date=structure.get('effective_date')
            )

            # 섹션 생성
            sections = []

            # 단락들을 섹션으로 그룹화
            if 'paragraphs' in content:
                current_section = DocumentSection(
                    id="section_1",
                    path=["root"],
                    heading="Document Content",
                    level=1,
                    blocks=[]
                )

                for i, para in enumerate(content['paragraphs']):
                    block = ContentBlock(
                        id=f"block_{i}",
                        type=ContentBlockType.PARAGRAPH,
                        page=1,
                        bbox=BoundingBox(0, 0, 100, 20, 1),  # 기본 좌표
                        content={"text": para['text']},
                        semantic=SemanticInfo(
                            keywords=['paragraph'],
                            confidence=0.9
                        )
                    )
                    current_section.blocks.append(block)

                sections.append(current_section)

            # 표 처리
            if 'tables' in content:
                for i, table in enumerate(content['tables']):
                    table_section = DocumentSection(
                        id=f"table_section_{i}",
                        path=["root", f"table_{i}"],
                        heading=f"Table {i+1}",
                        level=2,
                        blocks=[]
                    )

                    # 표 데이터 안전하게 처리
                    table_rows = []
                    if isinstance(table, dict):
                        if 'rows' in table and isinstance(table['rows'], list):
                            table_rows = table['rows']
                        elif 'data' in table and isinstance(table['data'], list):
                            table_rows = table['data']

                    row_count = len(table_rows) if table_rows else 0

                    table_block = ContentBlock(
                        id=f"table_{i}",
                        type=ContentBlockType.TABLE,
                        page=1,
                        bbox=BoundingBox(0, 0, 100, 50, 1),
                        content={"text": f"Table with {row_count} rows", "table_data": table_rows},
                        semantic=SemanticInfo(
                            keywords=['table'],
                            confidence=0.95
                        )
                    )
                    table_section.blocks.append(table_block)
                    sections.append(table_section)

            # 프로세스 흐름이 있으면 메타데이터에 추가
            if 'process_flows' in content and content['process_flows']:
                # 프로세스 흐름을 커스텀 필드에 추가
                metadata.source = content['process_flows']

            return DocJSON(
                version="2.0",
                doc_id=f"doc_{Path(file_name).stem}",
                metadata=metadata,
                sections=sections
            )

        except Exception as e:
            logger.error(f"DocJSON 변환 실패: {e}")
            return None